"""
Generates the InBetIO × Live xG Model integration guide as a polished .docx,
written for a non-technical audience (the project owner reads it before
sharing with the InBet dev team). Shorter and lighter than the v1 version.

Run:    python3 tools/build_integration_doc.py
Output: /tmp/InBetIO_Integration_Guide_WC2026.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x0F, 0x34, 0x60)
GOLD        = RGBColor(0xF5, 0xC5, 0x18)
INK         = RGBColor(0x1A, 0x1A, 0x2E)
MUTED       = RGBColor(0x71, 0x80, 0x96)
GREEN       = RGBColor(0x38, 0xA1, 0x69)
RED         = RGBColor(0xC5, 0x37, 0x37)
AMBER       = RGBColor(0xD6, 0x9E, 0x2E)
PURPLE      = RGBColor(0x80, 0x5A, 0xD5)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE  = RGBColor(0xEB, 0xF8, 0xFF)
LIGHT_GREEN = RGBColor(0xF0, 0xFF, 0xF4)
LIGHT_GOLD  = RGBColor(0xFF, 0xFB, 0xEB)
LIGHT_LILAC = RGBColor(0xFA, 0xF5, 0xFF)
LIGHT_GRAY  = RGBColor(0xF8, 0xFA, 0xFC)
TABLE_HDR   = NAVY

BOT_USERNAME = "@InBetWC2026_bot"
BOT_DEEPLINK = "https://t.me/InBetWC2026_bot?start={member_uuid}"

# ── Helpers ──────────────────────────────────────────────────────────────────
def hex(rgb):
    return str(rgb).upper()

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex(rgb))
    tcPr.append(shd)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'nil')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def set_para_border_left(para, rgb, sz='24'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), hex(rgb))
    pBdr.append(left)
    pPr.append(pBdr)

def set_run(run, size_pt, bold=False, color=None, italic=False, name='Calibri'):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_spacer(doc, height_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run('')
    r.font.size = Pt(height_pt)

def add_rule(doc, color=NAVY, sz='8'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex(color))
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_h1(doc, text):
    add_spacer(doc, 10)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_run(r, 18, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex(NAVY))
    pBdr.append(bottom)
    pPr.append(pBdr)
    add_spacer(doc, 4)

def add_h2(doc, text):
    add_spacer(doc, 4)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_run(r, 13, bold=True, color=NAVY)

def add_p(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    set_run(r, 11, italic=italic, color=color)
    return p

def add_bullet(doc, text, color=None, indent=0.2):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(indent)
    r = p.add_run(text)
    set_run(r, 11, color=color)

def add_box(doc, icon, title, body, bg, accent):
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)
    tbl.columns[0].width = Cm(1.0)
    tbl.columns[1].width = Cm(14.5)
    icon_cell, text_cell = tbl.rows[0].cells
    set_cell_bg(icon_cell, bg)
    set_cell_bg(text_cell, bg)
    p1 = icon_cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(8)
    r = p1.add_run(icon)
    set_run(r, 14)
    p2 = text_cell.paragraphs[0]
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after  = Pt(2)
    set_para_border_left(p2, accent, sz='24')
    if title:
        rt = p2.add_run(title + '\n')
        set_run(rt, 10, bold=True, color=accent)
    rb = p2.add_run(body)
    set_run(rb, 10, color=RGBColor(0x2D, 0x37, 0x48))
    add_spacer(doc, 6)

def add_table(doc, headers, rows, col_widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n)
    remove_table_borders(tbl)
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, TABLE_HDR)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(h)
        set_run(r, 9, bold=True, color=WHITE)
    for ri, row in enumerate(rows):
        bg = LIGHT_GRAY if ri % 2 == 1 else WHITE
        cells = tbl.rows[ri+1].cells
        for ci, val in enumerate(row):
            if isinstance(val, dict):
                bg_o  = val.get('bg', bg)
                text  = val.get('text', '')
                bold  = val.get('bold', False)
                color = val.get('color', RGBColor(0x2D, 0x37, 0x48))
            else:
                bg_o, text, bold, color = bg, str(val), False, RGBColor(0x2D, 0x37, 0x48)
            set_cell_bg(cells[ci], bg_o)
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            set_run(r, 10, bold=bold, color=color)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    add_spacer(doc, 8)

def add_link_row(doc, label, url, badge_color=NAVY):
    """One-line row: bold label · monospace url (kept on one para to copy easily)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.1)
    r1 = p.add_run('▸  ')
    set_run(r1, 10, bold=True, color=badge_color)
    r2 = p.add_run(label + '   ')
    set_run(r2, 10, bold=True, color=INK)
    r3 = p.add_run(url)
    set_run(r3, 9, color=NAVY, name='Courier New')

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex(RGBColor(0x1A, 0x1A, 0x2E)))
    pPr.append(shd)
    set_para_border_left(p, GOLD, sz='20')
    r = p.add_run(text)
    set_run(r, 9, name='Courier New', color=RGBColor(0xE2, 0xE8, 0xF0))

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── COVER (clean, no shading) ────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run('INTEGRATION GUIDE')
set_run(r, 10, bold=True, color=MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('InBetIO × Live xG Model')
set_run(r, 30, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run('For the FIFA World Cup 2026')
set_run(r, 14, color=MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run(f'13 May 2026   ·   Bot {BOT_USERNAME}   ·   Confidential')
set_run(r, 9, color=MUTED)

add_rule(doc)
add_spacer(doc, 10)

# ══ 1. AT A GLANCE ═══════════════════════════════════════════════════════════
add_h1(doc, '1.  At a glance')

add_p(doc,
    "We've built an AI football-prediction system that already runs live for "
    "the webpronos.com audience. For the 2026 World Cup, we're packaging it as "
    "a white-label product for InBetIO premium members under the brand "
    "\"Powered by InBetIO Live xG Model\".")

add_bullet(doc, "Two ready-to-embed widgets — one live match panel, one performance dashboard.")
add_bullet(doc, "A dedicated Telegram bot that pings premium members the moment a value pick is detected.")
add_bullet(doc, "Four languages out of the box: English, Spanish, European Portuguese, Brazilian Portuguese.")
add_bullet(doc, "Tournament window only: 11 June – 19 July 2026.")
add_bullet(doc, "InBet integration effort: roughly 4 hours total.")

add_box(doc, '🤝', 'Two-side split',
    'We host, update and maintain the entire system. InBetIO embeds two iframes, '
    'adds a button to open the Telegram bot, and exposes one small API endpoint '
    'we can query to check membership status. Everything else is on us.',
    LIGHT_GREEN, GREEN)

# ══ 2. HOW THE ALGORITHM WORKS ═══════════════════════════════════════════════
add_h1(doc, '2.  How the algorithm works')

add_p(doc,
    "Every 30 seconds, the system reads live match data from Sofascore, "
    "recalculates each team's Expected Goals (xG) based on the shots they've "
    "just taken, compares the result against current bookmaker odds, and — if "
    "the model finds a clear edge — surfaces a value pick. That pick goes "
    "instantly to the widgets and, in parallel, to the Telegram bot.")

# Insert pipeline image (centred)
add_spacer(doc, 4)
img_para = doc.add_paragraph()
img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_para.paragraph_format.space_before = Pt(2)
img_para.paragraph_format.space_after  = Pt(4)
img_run  = img_para.add_run()
img_run.add_picture('/tmp/pipeline_funnel.png', width=Inches(6.4))

# Caption
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run('Figure 1 · End-to-end pipeline from Sofascore to the member')
set_run(r, 9, italic=True, color=MUTED)

add_p(doc,
    "The whole flow takes 1–3 seconds per cycle. When a match ends, every pick "
    "is automatically settled (won / lost / push) so the dashboard's totals are "
    "always up to date.")

# ══ 3. THE TWO WIDGETS ═══════════════════════════════════════════════════════
add_h1(doc, '3.  The two widgets')

add_p(doc,
    "Both widgets are iframes — InBetIO drops one line of HTML wherever you "
    "want them on the members area. No JavaScript on your side.")

# 3.1 Per-match widget
add_h2(doc, 'Widget 1 · Per-Match')
add_p(doc,
    "A single panel that automatically shifts through five visual states as "
    "the match day unfolds: a live game, the post-match results (winning or "
    "losing version), a preview of the next fixture, or an off-day card on "
    "rest days. The same iframe URL covers all of them — the backend decides "
    "what to show.")

add_code(doc,
    '<iframe\n'
    '  src="https://embed.webpronos.com/widget/wc2026/current?lang=pt-pt&theme=dark"\n'
    '  width="100%" height="700"\n'
    '  frameborder="0" loading="lazy"\n'
    '  title="Live AI Picks — World Cup 2026">\n'
    '</iframe>')

add_p(doc, 'Query params: lang=en|es|pt-pt|pt-br · theme=dark|light · accent=#HEXCOLOR (optional)', italic=True, color=MUTED)

# Preview links (mock + production)
add_h2(doc, 'Preview links you can click NOW (mock data)')
add_p(doc,
    "These open in any browser and show exactly what each state looks like, "
    "using realistic-looking mock data so you can validate the design before "
    "the tournament starts.", italic=True, color=MUTED)

mock_base = 'https://livexgmodel-pt.fly.dev/widget/wc2026/current'
add_link_row(doc, 'LIVE match',              f'{mock_base}?mock=live&lang=pt-pt',         RED)
add_link_row(doc, 'Results · winning day',   f'{mock_base}?mock=results_win&lang=pt-pt',  GREEN)
add_link_row(doc, 'Results · losing day',    f'{mock_base}?mock=results_loss&lang=pt-pt', AMBER)
add_link_row(doc, 'Preview · next fixture',  f'{mock_base}?mock=preview&lang=pt-pt',      NAVY)
add_link_row(doc, 'Off-day rest card',       f'{mock_base}?mock=off_day&lang=pt-pt',      MUTED)

add_spacer(doc, 4)

# 3.2 Performance dashboard
add_h2(doc, 'Widget 2 · Performance Dashboard')
add_p(doc,
    "Tournament-wide proof of value: total picks, win rate, profit, ROI, "
    "equity curve, top 5 biggest winners, profit by market. Refreshes "
    "every 5 minutes.")

add_code(doc,
    '<iframe\n'
    '  src="https://embed.webpronos.com/widget/wc2026/performance?lang=pt-pt&theme=dark"\n'
    '  width="100%" height="560"\n'
    '  frameborder="0" loading="lazy"\n'
    '  title="InBetIO Live xG Model — Performance">\n'
    '</iframe>')

add_link_row(doc, 'Dashboard preview (mock data)',
    'https://livexgmodel-pt.fly.dev/widget/wc2026/performance?mock=1&lang=pt-pt', PURPLE)

add_box(doc, '🌐', 'Production URLs (used after the embed subdomain goes live)',
    'embed.webpronos.com is the public host. Until the CNAME is created in '
    'Cloudflare (Part 6), the same paths work on https://livexgmodel-pt.fly.dev '
    'with identical behaviour.',
    LIGHT_BLUE, RGBColor(0x31, 0x82, 0xCE))

# ══ 4. THE TELEGRAM BOT ══════════════════════════════════════════════════════
add_h1(doc, '4.  The Telegram bot')

add_p(doc,
    f"A dedicated InBetIO-branded Telegram bot ({BOT_USERNAME}) sends each new "
    "value pick to eligible members in real time. The binding is one-time per "
    "member and takes ~3 seconds.")

add_h2(doc, 'What the member does')
add_bullet(doc, 'Clicks a "Get Pick Alerts" button you place in the members area.')
add_bullet(doc, f'Telegram opens automatically pointing at {BOT_USERNAME} with their member ID embedded in the start URL.')
add_bullet(doc, 'Taps Start. The bot replies "✅ Account linked!" in their language.')
add_bullet(doc, 'From that moment, every new pick during the World Cup arrives as a Telegram message — until they /stop or their plan expires.')

add_h2(doc, 'What InBetIO places in the members area')
add_code(doc,
    f'<a href="{BOT_DEEPLINK.replace("{member_uuid}", "{{member_uuid}}")}"\n'
    f'   target="_blank" rel="noopener">\n'
    '   Get Pick Alerts\n'
    '</a>')
add_p(doc,
    'Replace {member_uuid} with the unique ID you assign to each member in your system. '
    'This ID is how we match a Telegram chat to a paying InBetIO account.',
    italic=True, color=MUTED)

add_h2(doc, 'Who actually receives picks')
add_table(doc,
    ['Plan status', 'Receives picks?'],
    [
        ['premium',  {'text':'✓  Yes', 'bold':True, 'color':GREEN}],
        ['trial',    {'text':'✓  Yes', 'bold':True, 'color':GREEN}],
        ['demo',     {'text':'✓  Yes', 'bold':True, 'color':GREEN}],
        ['inactive', {'text':'✗  No (paused — re-activates instantly if plan resumes)', 'bold':False, 'color':MUTED}],
        ['expired',  {'text':'✗  No', 'bold':False, 'color':MUTED}],
    ],
    col_widths=[2.0, 4.5])

add_box(doc, '📬', 'Every value pick, no throttling',
    'Paid members receive every value pick the algorithm surfaces — there is '
    'no artificial cap. In practice the model produces 1–4 picks per match, so '
    'this is well within reasonable notification volume. Members who prefer '
    'fewer pings can pause anytime with /stop and resume with /resume.',
    LIGHT_GREEN, GREEN)

add_h2(doc, 'Example message a member receives')
add_code(doc,
    '🏆  World Cup 2026 · Live Pick Alert\n\n'
    "🏴󠁧󠁢󠁥󠁮󠁧󠁿 England vs 🇧🇷 Brazil · 34'\n"
    'Over 2.5 Goals @ 1.85  ·  +8.2% xG edge\n\n'
    'Powered by InBetIO Live xG Model')

# ══ 5. WHAT INBETIO NEEDS TO DO ══════════════════════════════════════════════
add_h1(doc, '5.  What InBetIO needs to do')

add_p(doc, "Total estimated effort: ~4 hours, spread as follows.")

add_table(doc,
    ['#', 'Task', 'Effort', 'Required?'],
    [
        ['1', 'Paste 2 iframes on the members area (one for the live match panel, one for the performance dashboard).',  '~30 min', {'text':'Yes','bold':True,'color':GREEN}],
        ['2', f'Add a "Get Pick Alerts" button linking to the Telegram bot deep-link.',                                   '~30 min', {'text':'Yes','bold':True,'color':GREEN}],
        ['3', 'Expose a small GET endpoint we can call to fetch a member\'s plan status. Returns plan_status + locale.',  '~2 h',     {'text':'Yes','bold':True,'color':GREEN}],
        ['4', 'Generate a shared secret (e.g. via openssl rand -hex 32) and send it to us so we can authenticate the call above.', '~5 min', {'text':'Yes','bold':True,'color':GREEN}],
        ['5', 'Add a CNAME record in Cloudflare for embed.webpronos.com → livexgmodel-pt.fly.dev.',                       '~5 min', {'text':'Yes','bold':True,'color':GREEN}],
        ['6', 'Call our webhook when a member\'s plan status changes (instant sync; without it we re-check every 24 h).', '~1 h',   {'text':'Optional','color':MUTED}],
        ['7', 'Call our unlink endpoint when a member deletes their account.',                                            '~30 min', {'text':'Optional','color':MUTED}],
    ],
    col_widths=[0.3, 3.8, 0.8, 1.6])

add_box(doc, '🔑', 'The shared secret',
    'A single string (32 chars or more) that proves any HTTP call between our '
    'two systems is genuine. You generate it once, share it with us once, and '
    'we both store it on our backends — never exposed to browsers, never logged.',
    LIGHT_GREEN, GREEN)

# ══ 6. BEHIND THE SCENES ═════════════════════════════════════════════════════
add_h1(doc, '6.  Behind the scenes')

add_h2(doc, 'Where we read live data from')
add_p(doc,
    "All match data — events, scores, shot maps, xG, live momentum, upcoming "
    "fixtures — comes from Sofascore's public API (api.sofascore.com). "
    "Sofascore is the industry-standard live-data source used by most "
    "betting and analytics products you've seen.")
add_link_row(doc, 'Reference', 'https://api.sofascore.com/api/v1/', NAVY)
add_p(doc,
    "Our backend wraps the Sofascore API with a fingerprint-rotating client to "
    "stay stable under regional rate limits. We do not store or resell "
    "Sofascore data — we consume it live, derive picks, and discard the rest.",
    italic=True, color=MUTED)

add_h2(doc, 'Where everything runs')
add_p(doc,
    "Hosting is on Fly.io — a developer-focused cloud that lets a single app "
    "run from any of 14+ global regions. Fly.io was chosen because it makes "
    "the next paragraph possible.")
add_link_row(doc, 'Provider', 'https://fly.io/', NAVY)

add_h2(doc, 'Resilience — when Sofascore blocks us')
add_p(doc,
    "Occasionally Sofascore rate-limits or blocks the IP range of whichever "
    "region our app is running in. When that happens:")
add_bullet(doc, 'A watchdog (GitHub Actions, hourly) detects two consecutive failed health-checks.')
add_bullet(doc, 'It spawns disposable test machines in 14 candidate regions in parallel.')
add_bullet(doc, 'The first region that successfully reaches Sofascore wins.')
add_bullet(doc, 'The production volume is forked to that region and traffic moves there — usually within ~5 minutes.')
add_bullet(doc, 'You and the admin get a Telegram alert confirming the migration.')
add_box(doc, '⚙️', 'Why this matters for InBetIO',
    'The system effectively heals itself. Even when Sofascore blocks a region, '
    'downtime is bounded to under 2 hours — InBetIO members keep getting picks.',
    LIGHT_LILAC, PURPLE)

# ══ 7. PRE-LAUNCH PLAN ═══════════════════════════════════════════════════════
add_h1(doc, '7.  Pre-launch plan')

add_bullet(doc, 'Now → end-May: InBet completes Tasks 1–5 from §5. We complete final UI polish on our side.')
add_bullet(doc, 'Early June: Joint dry-run using one real, live football match (any league). We confirm widgets render correctly, Telegram fan-out works, status sync behaves.')
add_bullet(doc, 'Tuesday 9 June: Full freeze. No more code changes. We monitor production.')
add_bullet(doc, 'Thursday 11 June: Tournament opener — South Africa vs Mexico. Both teams on standby; any bug found triggers a hotfix during the match.')

# ── FOOTER ───────────────────────────────────────────────────────────────────
add_spacer(doc, 14)
add_rule(doc, MUTED, sz='4')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r1 = p.add_run('InBetIO × Live xG Model · World Cup 2026   ·   ')
set_run(r1, 9, color=MUTED)
r2 = p.add_run('Questions? Add your contact details here before sharing.')
set_run(r2, 9, bold=True, color=NAVY)
r3 = p.add_run('   ·   🔒 Confidential')
set_run(r3, 9, color=MUTED)

# ── SAVE ─────────────────────────────────────────────────────────────────────
out = '/tmp/InBetIO_Integration_Guide_WC2026.docx'
doc.save(out)
print(f'✅  Saved: {out}')
